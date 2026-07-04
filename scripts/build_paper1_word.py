from __future__ import annotations

import re
import shutil
import subprocess
import sys
import hashlib
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\paper_MedIA Vol. 107–113\failure_region_reliability")
TEX = ROOT / "manuscript" / "pdf_draft" / "paper1_sfrm_audit_draft.tex"
FIGURES = ROOT / "figures" / "paper1"
ATTACHMENT = Path(r"C:\Users\LIU\.openclaw-autoclaw\fig1_seedream_output.jpg")
OUT_DIR = ROOT / "manuscript" / "word_draft"
WORK_DIR = OUT_DIR / "build"
DOCX = OUT_DIR / "paper1_sfrm_audit_draft_word.docx"
PANDOC = Path(r"C:\Users\LIU\AppData\Local\Pandoc\pandoc.exe")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if len(widths_dxa) != len(table.columns):
        raise ValueError(f"Geometry has {len(widths_dxa)} widths for {len(table.columns)} columns")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)


def preprocess_tex() -> Path:
    if not TEX.exists():
        raise FileNotFoundError(TEX)
    if not ATTACHMENT.exists():
        raise FileNotFoundError(ATTACHMENT)
    if not PANDOC.exists():
        raise FileNotFoundError(PANDOC)

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    replacement = WORK_DIR / "figure2_sfrm_framework_replacement.jpg"
    shutil.copy2(ATTACHMENT, replacement)

    text = TEX.read_text(encoding="utf-8")
    bib_keys = re.findall(r"\\bibitem\{([^{}]+)\}", text)
    bib_index = {key: idx + 1 for idx, key in enumerate(bib_keys)}

    def replace_citation(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        numbers = [str(bib_index[key]) for key in keys if key in bib_index]
        if len(numbers) != len(keys):
            missing = [key for key in keys if key not in bib_index]
            raise RuntimeError(f"Citation keys missing from bibliography: {missing}")
        return "[" + ", ".join(numbers) + "]"

    # Pandoc does not connect LaTeX's handwritten thebibliography entries to
    # citation commands. Materialize the numeric citations and numbered list so
    # the Word manuscript retains a complete, internally consistent reference trail.
    text = re.sub(r"\\cite\{([^{}]+)\}", replace_citation, text)
    text = text.replace("\\begin{thebibliography}{99}", "\\section*{References}\n\\begin{enumerate}")
    text = re.sub(r"\\bibitem\{[^{}]+\}", r"\\item", text)
    text = text.replace("\\end{thebibliography}", "\\end{enumerate}")
    # Word cannot embed PDF figures reliably. Every project figure has a PNG twin.
    text = re.sub(r"(\\includegraphics(?:\[[^\]]*\])?\{)([^{}]+)\.pdf(\})", r"\1\2.png\3", text)
    # Figure 2 in the compiled manuscript is the framework figure (the correlation
    # heatmap appears first and is Figure 1).
    text = text.replace("{fig1_sfrm_framework.png}", "{figure2_sfrm_framework_replacement.jpg}")
    temp_tex = WORK_DIR / "paper1_word_source.tex"
    temp_tex.write_text(text, encoding="utf-8")
    return temp_tex


def run_pandoc(temp_tex: Path) -> Path:
    raw_docx = WORK_DIR / "paper1_raw.docx"
    resource_path = f"{FIGURES};{WORK_DIR};{TEX.parent}"
    cmd = [
        str(PANDOC),
        str(temp_tex),
        "--from=latex",
        "--to=docx",
        "--standalone",
        f"--resource-path={resource_path}",
        "--output",
        str(raw_docx),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
    if result.returncode != 0:
        raise RuntimeError(f"Pandoc failed:\n{result.stdout}\n{result.stderr}")
    if not raw_docx.exists() or raw_docx.stat().st_size == 0:
        raise RuntimeError("Pandoc did not create a usable DOCX")
    return raw_docx


def style_document(raw_docx: Path) -> None:
    doc = Document(raw_docx)

    # Narrative-proposal preset resolved for an academic manuscript, with named
    # overrides for A4 paper and Times New Roman journal typography.
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.35)
        section.right_margin = Cm(2.35)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    heading_specs = {
        "Title": (17, 0, 10, RGBColor(0, 0, 0)),
        "Heading 1": (14, 12, 6, RGBColor(0, 0, 0)),
        "Heading 2": (11.5, 9, 4, RGBColor(0, 0, 0)),
        "Heading 3": (10.5, 7, 3, RGBColor(0, 0, 0)),
    }
    for style_name, (size, before, after, color) in heading_specs.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Title"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for caption_style_name in ("Caption", "Image Caption", "Table Caption"):
        if caption_style_name not in doc.styles:
            continue
        cap = doc.styles[caption_style_name]
        cap.font.name = "Times New Roman"
        cap.font.size = Pt(9)
        cap.font.color.rgb = RGBColor(0, 0, 0)
        cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(7)
        cap.paragraph_format.keep_with_next = caption_style_name == "Table Caption"

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name in {"Author", "Date"}:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(3)
        elif style_name == "Abstract":
            para.paragraph_format.left_indent = Cm(0.45)
            para.paragraph_format.right_indent = Cm(0.45)
            para.paragraph_format.space_after = Pt(5)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif style_name in {"Caption", "Image Caption", "Table Caption"}:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Keep figure paragraphs centered and prevent captions from separating.
        if para._p.xpath(".//w:drawing"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.keep_with_next = True

        for run in para.runs:
            set_run_font(run, "Times New Roman")

    figure_number = 0
    table_number = 0
    figure_alt_texts: list[str] = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Image Caption":
            figure_number += 1
            if para.runs:
                para.runs[0].text = f"Figure {figure_number}. " + para.runs[0].text
                para.runs[0].bold = True
            figure_alt_texts.append(para.text[:500])
        elif style_name == "Table Caption":
            table_number += 1
            if para.runs:
                para.runs[0].text = f"Table {table_number}. " + para.runs[0].text
                para.runs[0].bold = True

    if len(doc.inline_shapes) != len(figure_alt_texts):
        raise RuntimeError("Figure and caption counts do not match")
    for index, (shape, alt_text) in enumerate(zip(doc.inline_shapes, figure_alt_texts), start=1):
        shape._inline.docPr.set("title", f"Figure {index}")
        shape._inline.docPr.set("descr", alt_text)

    # Reference-list items use a compact academic rhythm while retaining true
    # Word numbering generated by Pandoc.
    reference_heading_seen = False
    for para in doc.paragraphs:
        if para.text.strip() == "References":
            reference_heading_seen = True
            continue
        if reference_heading_seen:
            if para.style and para.style.name.startswith("Heading"):
                reference_heading_seen = False
                continue
            para.paragraph_format.space_after = Pt(2.5)
            para.paragraph_format.line_spacing = 1.0
            for run in para.runs:
                set_run_font(run, "Times New Roman", 9)

    table_widths = [
        [1200, 1300, 900, 900, 1300, 1200, 2440],
        [1400, 1500, 700, 1000, 1000, 1000, 1400, 1240],
        [1300, 1300, 650, 900, 1600, 900, 1300, 1290],
        [1400, 1500, 700, 1000, 1000, 1000, 1400, 1240],
        [2000, 2000, 1000, 1500, 1000, 1740],
    ]
    if len(doc.tables) != len(table_widths):
        raise RuntimeError(f"Unexpected table count: {len(doc.tables)}")

    for table, widths in zip(doc.tables, table_widths):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_geometry(table, widths)
        if table.rows:
            set_repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    set_cell_shading(cell, "E9EDF2")
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(1.5)
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        set_run_font(run, "Times New Roman", 8.5)
                        if row_index == 0:
                            run.bold = True

    # Quiet running footer with a dynamic PAGE field.
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_char1, instr_text, fld_char2])
        set_run_font(run, "Times New Roman", 9)

    # Core properties are intentionally generic for an internal review draft.
    doc.core_properties.title = "Beyond Global Uncertainty: Scale Heterogeneity in Medical Image Segmentation Failure Detection"
    doc.core_properties.subject = "SFRM manuscript draft"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)


def verify_structure() -> None:
    doc = Document(DOCX)
    images = len(doc.inline_shapes)
    tables = len(doc.tables)
    equations = len(doc._element.xpath(".//m:oMath | .//m:oMathPara"))
    text = "\n".join(p.text for p in doc.paragraphs)
    required = [
        "Beyond Global Uncertainty",
        "Spatial Failure-Region Modeling",
        "Primary low-boundary-quality results",
        "Data and Code Availability",
        "References",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Required manuscript content missing from DOCX: {missing}")
    if images < 9:
        raise RuntimeError(f"Expected at least 9 figures, found {images}")
    if tables < 4:
        raise RuntimeError(f"Expected at least 4 tables, found {tables}")
    if equations < 8:
        raise RuntimeError(f"Expected at least 8 equations, found {equations}")
    if "[1]" not in text or "[29, 30]" not in text:
        raise RuntimeError("Numbered citation materialization failed")
    if text.count("Figure ") < 9 or text.count("Table ") < 5:
        raise RuntimeError("Figure/table caption numbering failed")

    source_hash = hashlib.sha256(ATTACHMENT.read_bytes()).hexdigest()
    with zipfile.ZipFile(DOCX) as archive:
        media_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
    if source_hash not in media_hashes:
        raise RuntimeError("The requested Figure 2 attachment was not embedded byte-for-byte")
    print(f"DOCX={DOCX}")
    print(f"paragraphs={len(doc.paragraphs)} tables={tables} images={images} equations={equations}")
    print(f"bytes={DOCX.stat().st_size}")


def main() -> None:
    temp_tex = preprocess_tex()
    raw_docx = run_pandoc(temp_tex)
    style_document(raw_docx)
    verify_structure()


if __name__ == "__main__":
    main()
